import streamlit as st
from google import genai
from google.genai import types
import os
import io
from docx import Document  # Ensure you ran: pip install python-docx
from docx.shared import Pt
from dotenv import load_dotenv

# 1. BOOT SYSTEM & API
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

# Initialize Client with 2026 Preview Model support
if api_key:
    client = genai.Client(
        api_key=api_key,
        http_options=types.HttpOptions(api_version='v1beta')
    )
else:
    st.error("🔑 API Key missing! Check your .env file.")
    st.stop()

# 2. WORD DOCUMENT GENERATOR
def generate_word_report(shortlist):
    doc = Document()
    doc.add_heading('🛡️ My Warrior Job Shortlist', 0)
    
    for job in shortlist:
        doc.add_heading('Opportunity Analysis', level=1)
        
        # Original Job Snippet
        doc.add_heading('Job Snippet:', level=2)
        doc.add_paragraph(job['Job_Content'])
        
        # AI Verdict and Email Plan
        doc.add_heading('AI Strategy & Verdict:', level=2)
        p = doc.add_paragraph(job['AI_Verdict'])
        p.style.font.size = Pt(11)
        
        doc.add_page_break() # Fresh page for every job
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 3. SESSION MEMORY
if 'job_queue' not in st.session_state: st.session_state.job_queue = []
if 'current_idx' not in st.session_state: st.session_state.current_idx = 0
if 'shortlist' not in st.session_state: st.session_state.shortlist = []

# 4. UI LAYOUT
st.set_page_config(page_title="Warrior Job Catcher", layout="wide", page_icon="🛡️")

st.title("🛡️ Warrior Job Catcher v4.0")
st.caption("Pune Regional Hub | Gemini 3 Flash | Word Export Enabled")

# SIDEBAR: THE INPUT HUB
with st.sidebar:
    st.header("📥 Job Feed")
    st.write("Target: 8.4 CGPA | Python Internships")
    
    raw_input = st.text_area("Paste Jobs (Separator: '---')", height=300, 
                             placeholder="Paste text here...")
    
    if st.button("🚀 Load Jobs", use_container_width=True):
        jobs = [j.strip() for j in raw_input.split("---") if len(j.strip()) > 50]
        st.session_state.job_queue = jobs
        st.session_state.current_idx = 0
        st.rerun()

    st.divider()
    
    # THE DOWNLOAD HUB
    if st.session_state.shortlist:
        st.header("📋 Export Plan")
        word_file = generate_word_report(st.session_state.shortlist)
        st.download_button(
            label="📥 Download as Word (.docx)",
            data=word_file,
            file_name="Pune_Internship_Battle_Plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# 5. THE ENGINE
# --- UPDATED ANALYZER ENGINE (Overload-Proof) ---
if st.session_state.job_queue and st.session_state.current_idx < len(st.session_state.job_queue):
    # Only take a snippet to prevent overloading the model
    current_job_text = st.session_state.job_queue[st.session_state.current_idx][:4000]
    
    col_desc, col_ai = st.columns([3, 2])
    
    with col_desc:
        st.info(f"Reviewing Job {st.session_state.current_idx + 1}")
        st.markdown(f"```text\n{current_job_text[:1000]}...\n```") # UI only shows snippet

    with col_ai:
        st.subheader("🤖 Gemini 3 Verdict")
        try:
            with st.spinner("Analyzing..."):
                # Simpler prompt = faster, more stable response
                response = client.models.generate_content(
                    model='gemini-3-flash-preview',
                    contents=f"Analyze this Pune Python Intern job for an 8.4 CGPA student. Give Match Score, Verdict, and a quick Email Draft: {current_job_text}"
                )
                analysis = response.text
                st.write(analysis)
            
            # Action Buttons
            st.divider()
            c1, c2 = st.columns(2)
            if c1.button("✅ ACCEPT", use_container_width=True):
                st.session_state.shortlist.append({
                    "Job_Content": current_job_text[:200],
                    "AI_Verdict": analysis
                })
                st.session_state.current_idx += 1
                st.rerun()
            
            if c2.button("❌ REJECT", use_container_width=True):
                st.session_state.current_idx += 1
                st.rerun()

        except Exception as e:
            if "429" in str(e) or "overloaded" in str(e).lower():
                st.warning("🚦 Model is a bit busy. Wait 5 seconds and click Refresh.")
                if st.button("🔄 Try Again"):
                    st.rerun()
            else:
                st.error(f"Engine Error: {e}")

elif st.session_state.job_queue:
    st.balloons()
    st.success("🎯 Queue Cleared! Your battle plan is ready in the sidebar.")
    if st.button("Restart Queue"):
        st.session_state.job_queue = []
        st.rerun()
else:
    st.info("👋 System Standby. Feed the sidebar with job descriptions to begin.")